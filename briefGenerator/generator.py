# 将控制台内容重定向至文件
"""
使用os.system(command)等模块执行系统命令时，返回值为命令执行结果，命令执行成功返回Ture，否则返回False；
若要得命令本身返回的内容，需要将命令输出至控制台的内容写到文件中，即将标准输出由控制台重定向至文件。
"""
import sys
import os
import time
from docx import Document

from aviation import aviast
from OceanBrief import generator as Oceangenerator

# 自定义目标文件夹和目标文件名
filepath = "E:/code/projects/weatherPlatform/templates"
filename = "chongDX.html"
fullname = os.path.join(filepath, filename)

# 备份默认的标准输出（输出值控制台）
standard_output = sys.stdout

# 将标准输出重定向至文件
"""
此处实质为向文件写入内容。
通常，向文件写入内容的步骤为：打开文件-->写入内容-->关闭文件
此处步骤与上述步骤相同，只不过“写”的方式发生了改变。
一般的文件写入为由人通过键盘键入内容或者copy内容，此处为由解释器向文件写入内容。
标准输出是解释器将内容写到（输出到）控制台，我们可以在控制台看到内容，此处这是解释器将内容
写入（输出到）了文件，我们可以在文件中找到内容。
"""

sys.stdout = open(fullname, "w+", encoding='utf-8')
print("爬虫开始")
sys.stdout = open(fullname, "a+", encoding='utf-8')
# 写入内容。已经将标注输出更改为输出至文件，所以执行命令后，会将原来输出至控制台的内容输出至文件。


def getcur():
    sep = '/'
    if os.getcwd().find(sep) == -1:
        sep = '\\'
    return os.getcwd().split(sep)[-1]

oceanptah = "OceanBrief"
aviationpath = "aviation"
trafficpath = "trafficbriefing"
weatherbriefpath = "weatherbriefing"

cur = time.strftime('%Y{y}%m{m}%d{d}', time.localtime(time.time())).format(y='年', m='月', d='日')

#第一部分完成爬取、数据存储
#第二部分完成简报生成
#简报命名 2020年05月12日简报（cur+"简报"）

#第一部分
#气象概况
if getcur() != weatherbriefpath:
    os.chdir(weatherbriefpath)
os.chdir("./weatherbriefingspider")
os.system("python start.py")
os.chdir("../")
os.chdir("../")
print("============weatherbrief ok================<br>")
sys.stdout = open(fullname, "a+", encoding='utf-8')




#do something
#...

#公路
if getcur() != trafficpath:
    os.chdir(trafficpath)
os.chdir("./trafficBriefing")
os.system("python start.py")
os.chdir("../")
os.chdir("../")
print("============traffic ok================<br>")
sys.stdout = open(fullname, "a+", encoding='utf-8')




#do something
#...


#水路
if getcur() != oceanptah:
    os.chdir(oceanptah)
#do something
#...
os.system("python start.py")
os.chdir("../")
print("============ocean ok================<br>")
sys.stdout = open(fullname, "a+",encoding='utf-8')





#航空
if getcur() != aviationpath:
    os.chdir(aviationpath)
#do something
#...
os.system('scrapy crawl airport')
os.chdir("../")
print("============aviation ok================")
sys.stdout = open(fullname, "a+" ,encoding='utf-8')




#生成简报
#概述
#公路
#水路
#民航
#公路附件
#水路附件
#民航附件
path = "./"+cur+"简报.docx"
if os.path.exists(path):
    #print(1)
    os.remove(path)

document = Document()
document.save(path)

#概述
os.chdir(weatherbriefpath)
os.chdir("./briefGenerator")
os.system("python MainGenerator.py")
os.chdir("../")
os.chdir("../")
print("============概述 ok================")
sys.stdout = open(fullname, "a+", encoding='utf-8')

# 写入内容。已经将标注输出更改为输出至文件，所以执行命令后，会将原来输出至控制台的内容输出至文件。


#公路

os.chdir(trafficpath)
os.chdir("./briefGenerator")
os.system("python MainGenerator.py")
os.chdir("../")
os.chdir("../")
print("============公路 ok================")
sys.stdout = open(fullname, "a+", encoding='utf-8')

# 写入内容。已经将标注输出更改为输出至文件，所以执行命令后，会将原来输出至控制台的内容输出至文件。



#水路
document = Document(path)
Oceangenerator.firstpart(document)
document.save(path)
print("============水路 ok================")
sys.stdout = open(fullname, "a+", encoding='utf-8')



# 关闭文件
sys.stdout.close()
#民航
document = Document(path)
aviast.firstpart(document)
document.save(path)
print("============民航 ok================")
sys.stdout = open(fullname, "a+", encoding='utf-8')

# 写入内容。已经将标注输出更改为输出至文件，所以执行命令后，会将原来输出至控制台的内容输出至文件。



#公路附件
os.chdir(trafficpath)
os.chdir("./briefGenerator")
os.system("python AddictionGenerator.py")
os.chdir("../")
os.chdir("../")
print("============公路附件 ok================")
#水路附件
document = Document(path)
Oceangenerator.secondpart(document)
document.save(path)
print("============水路附件 ok================")
#民航附件
document = Document(path)
aviast.secondpart(document)
document.save(path)
print("============民航附件 ok================")
# 关闭文件
sys.stdout.close()

# 恢复默认标准输出
sys.stdout = standard_output

