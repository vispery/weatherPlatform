from docx import Document
import os
import time

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn

def read_weatherforecast():
	forecastpath = os.path.abspath(os.path.dirname(__file__))+"\\forecast.txt"
	weatherforecast = ""
	with open(forecastpath,mode = "r",encoding="utf-8") as f:
		for line in f:
			weatherforecast = weatherforecast+line
		return weatherforecast


def firstpart(document):
	cur = time.strftime('%Y{y}%m{m}%d{d}', time.localtime(time.time())).format(y='年', m='月', d='日')
	#path = "../"+cur+"简报.docx"
	'''
	try:
		document = Document(path)
	except:
		document = Document()
	'''
	para = document.add_paragraph()
	run = para.add_run("")
	run.font.name = u'仿宋_GB2312'
	run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run.font.bold = True
	run.font.color.rgb = RGBColor(0, 0, 0)

	pic = document.add_paragraph()
	picpath = os.path.abspath(os.path.dirname(__file__))
	run = pic.add_run("")
	run.add_picture(picpath + '/OceanBrief/img/' + cur + 'wind.png', width=Inches(2.90))
	run.add_picture(picpath + '/OceanBrief/img/' + cur + 'map.png', width=Inches(2.90))
	forecast = read_weatherforecast()
	if(forecast[-1]==';'):
		forecast = forecast[0:-1]
	para = document.add_paragraph()
	run = para.add_run("海域。"+forecast + "。")
	run.font.name = u'仿宋_GB2312'
	run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run.font.size = Pt(12)
	run.font.color.rgb = RGBColor(0, 0, 0)
	document.add_page_break()



def secondpart(document):
	cur = time.strftime('%Y{y}%m{m}%d{d}', time.localtime(time.time())).format(y='年', m='月', d='日')
	#path = "../"+cur+"简报.docx"
	'''
	try:
		document = Document(path)
	except:
		document = Document()
	'''
	para = document.add_paragraph()
	run = para.add_run("五、全海域风、浪、雾预报：")
	run.font.bold = True
	run.font.name = u'黑体'
	run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
	run.font.size = Pt(16)

	para = document.add_paragraph()
	forecast = read_weatherforecast()
	forecastlist = forecast.split(';')
	for i in range(len(forecastlist)):
		if(len(forecastlist[i])==0):
			break
		run = para.add_run(str(i+1)+'.'+forecastlist[i]+'\n')
		run.font.name = u'仿宋_GB2312'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
		run.font.size = Pt(14)

