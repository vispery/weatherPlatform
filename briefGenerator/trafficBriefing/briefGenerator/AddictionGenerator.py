# -*- coding: utf-8 -*-
import time
import json
import os
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def num_to_char(num):
    """数字转中文"""
    num=str(num)
    new_str=""
    num_dict={"0":u"零","1":u"一","2":u"二","3":u"三","4":u"四","5":u"五","6":u"六","7":u"七","8":u"八","9":u"九"}
    listnum=list(num)
    # print(listnum)
    shu=[]
    for i in listnum:
        # print(num_dict[i])
        shu.append(num_dict[i])
    new_str="".join(shu)
    # print(new_str)
    return new_str

def add_detail(influenced,document,seq):
    sub_title=influenced[0]
    run = document.add_paragraph().add_run(num_to_char(seq)+'、 '+sub_title)
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # 三号
    run.font.size = Pt(16)
    run.font.bold = True
    snow_list = influenced[1:]
    for i,each in enumerate(snow_list):
        each = str(i+1)+'. '+each
        run = document.add_paragraph().add_run(each)
        run.font.name = u'仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run.font.size = Pt(14)

p=os.path.dirname(__file__)
p = os.path.join(p,os.path.pardir)
raw = {}
with open('{}/trafficBriefing/trafficBriefing/brief.json'.format(p), 'r', encoding='utf-8') as fObj:
    raw_list = json.load(fObj)
    raw = raw_list[-1]

date = time.strftime('%Y{y}%m{m}%d{d}',time.localtime(time.time())).format(y='年',m='月',d='日')
path=os.path.abspath(os.path.join(os.getcwd(), "../.."))+"\\"+date+"简报.docx"
document=Document(path)


strs = "陆域。"
content = ""


lenth = len(raw['snow_influenced'])
if lenth>1:
    strs+=raw['snow_influenced'][0][:-1]
    strs+=raw['snow_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['fog_influenced'])
if lenth>1:
    strs+='，'
    strs+=raw['fog_influenced'][0][:-1]
    strs+=raw['fog_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['rain_influenced'])
if lenth>1:
    strs+='，'
    strs+=raw['rain_influenced'][0][:-1]
    strs+=raw['rain_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['thunder_influenced'])
if lenth>1:
    strs+='，'
    strs+=raw['thunder_influenced'][0][:-1]
    strs+=raw['thunder_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['other_influenced'])
if lenth>0:
    for influenced in raw['other_influenced']:
        strs+='，'
        strs+=influenced[0][:-1]
        strs+=influenced[1]
        strs+="等" + str(len(influenced) - 1) + "条"


strs+='。'


p2_1 = document.add_paragraph().add_run("附件")
p2_1.font.name = u'方正小标宋简体'
p2_1._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
# 四号
p2_1.font.size = Pt(14)

title = "受天气影响的主要路段、海域、机场"
temp = document.add_paragraph()
temp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2_2 = temp.add_run(title)
p2_2.font.name = u'方正小标宋简体'
p2_2._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
# 二号
p2_2.font.size = Pt(22)
p2_2.font.bold = True

seq=1


lenth = len(raw['snow_influenced'])
if lenth>1:
    add_detail(raw['snow_influenced'],document,seq)
    seq+=1


lenth = len(raw['fog_influenced'])
if lenth>1:
    add_detail(raw['fog_influenced'],document,seq)
    seq+=1

lenth = len(raw['rain_influenced'])
if lenth>1:
    add_detail(raw['rain_influenced'],document,seq)
    seq+=1

lenth = len(raw['thunder_influenced'])
if lenth>1:
    add_detail(raw['thunder_influenced'],document,seq)
    seq+=1

lenth = len(raw['other_influenced'])
if lenth>0:
    for each in raw['other_influenced']:
        add_detail(each,document,seq)
        seq+=1


document.save(path)
