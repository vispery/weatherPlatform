# -*- coding: utf-8 -*-
import time
import json
import jieba
import re
import jieba
import time
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Length
key = ['暴雨','大雨','暴雪', '大雪', '沙尘暴']
num1=1
num2=1
num3=1
keyrain = ['暴雨','中到大雨','大雨','大到暴雨']
keysnow = ['暴雪','中到大雪','大雪','大到暴雪','中雪']
brief_flag = False
def change(line,str):
    part1 = re.split(r"[；，。（]", line)
    for tmp in part1:
        flag = 0
        word_list = jieba.lcut(tmp)
        for keyword in word_list:
            if keyword in key:
                global brief_flag
                brief_flag = True
                flag = 1;
        if flag == 1:
            str += '，'
            str += tmp
    return str

def change2(day,snow,rain,wind):
    global num1,num2,num3
    part1 = re.split(r"[，。]", day)
    #print(part1)
    for tmp in part1:
        flag = 0
        word_list = jieba.lcut(tmp)
        #print(word_list)
        for keyword in word_list:
            if keyword in keyrain:
                flag=1
            elif keyword in keysnow:
                flag=2
            elif keyword=='风':
                flag=3
        if flag == 1:
            rain +='（'+str(num1)+'）'+tmp+'。'
            num1+=1
        if flag == 2:
            snow +='（'+str(num2)+'）'+tmp+'。'
            num2+=1
        if flag == 3:
            wind +='（'+str(num3)+'）'+tmp+'。'
            num3+=1
    return rain,snow,wind


publish ="    "+time.strftime('%d', time.localtime(time.time())) + "日"
with open('../weatherBriefingSpider/weatherBriefingSpider/brief.json', 'r', encoding='utf-8') as fObj:
    raw_list = json.load(fObj)
    raw = raw_list[0]
    #完成发布部分内容
    line1 = raw['brief_detail']
    publish = change(line1, publish)
    publish += "。未来两至三日"
    line1 = raw['day1_detail']
    publish = change(line1, publish)
    publish = change(raw['day2_detail'], publish)
    publish = change(raw['day3_detail'], publish)
    if brief_flag:
        publish += "，可能对交通运输产生影响。"
    else:
        publish += "，没有特别恶劣天气。"
    #print(publish)

    #重要天气预报
    import_weather =""
    import_weather+=''.join(raw['key_point_title'])+'\n'
    import_weather+="    "+' '.join(raw['key_point_detail'])
    #print(import_weather)
  
    
date = time.strftime('%Y{y}%m{m}%d{d}',time.localtime(time.time())).format(y='年',m='月',d='日')
p=os.path.abspath(os.path.join(os.getcwd(), "../.."))+"\\"+date+"简报.docx"
document=Document(p)

# title
p0 = document.add_paragraph()
run=p0.add_run("                   每日交通气象预报信息") 
run.font.name = u'方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
run.font.size = Pt(22)

p1 = document.add_paragraph()
run=p1.add_run(publish) 
run.font.name = u'仿宋_GB2312'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
run.font.size = Pt(14)
p1.line_spacing_rule = WD_LINE_SPACING.MULTIPLE #多倍行距
p1.paragraph_format.line_spacing = 1.25 # 1.25倍行间距

document.save(p)