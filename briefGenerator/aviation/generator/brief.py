import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

#from aviation.generator.badlist import badlist
#from aviation.generator.image import getmap, processmap
from aviation.generator.badlist import badlist
from aviation.generator.image import getmap, processmap


def firstpart(document):
    '''
    filename = 'weatherbrief.docx'
    try:
        document = Document(filename)
    except PackageNotFoundError:
        document = Document()
    '''

    blis = badlist()

    # 概况
    getmap()
    processmap()
    img = document.add_paragraph().add_run()
    img.add_picture('airportmapprocessed.png', width=Cm(16.15))

    briefpara = document.add_paragraph()
    briefpara.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    briefpara.paragraph_format.left_indent = Pt(-80)
    briefpara.paragraph_format.right_indent = Pt(-80)

    kongyu = briefpara.add_run('空域。')
    kongyu.font.name = u'黑体'
    kongyu._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    kongyu.font.size = Pt(12)
    kongyu.font.color.rgb = RGBColor(0, 0, 0)

    run = briefpara.add_run('受不良天气影响的机场主要有：'+blis['badsetstr'])
    run.font.name = u'仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

    os.remove('airportmapprocessed.png')


def secondpart(document):
    '''
    filename = 'weatherbrief.docx'
    try:
        document = Document(filename)
    except PackageNotFoundError:
        document = Document()
    '''

    blis = badlist()

    # 附件
    title = document.add_paragraph().add_run('六、全国航空气象预报：')
    title.font.name = u'黑体'
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    title.font.size = Pt(16)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    numlist = document.styles['List Number']
    numlist.font.name = u'仿宋_GB2312'
    numlist._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    numlist.font.size = Pt(14)
    numlist.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for item in blis['badstrlist']:
        run = document.add_paragraph(style=numlist).add_run(item)
        run.font.name = u'仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
